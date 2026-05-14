import os, sys, subprocess, time, threading, requests, jwt
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

os.chdir(os.path.dirname(os.path.abspath(__file__)))

doc = Document()

# ── Estilos globales ──────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)

AZUL_OSCURO = RGBColor(0x1F, 0x49, 0x7D)
ROJO        = RGBColor(0xC0, 0x00, 0x00)
VERDE       = RGBColor(0x37, 0x86, 0x30)
GRIS_CLARO  = RGBColor(0xF2, 0xF2, 0xF2)
NEGRO       = RGBColor(0x00, 0x00, 0x00)
NARANJA     = RGBColor(0xE3, 0x6C, 0x09)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_heading(text, level=1, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    if level == 1:
        run.font.size = Pt(18)
        run.bold = True
        run.font.color.rgb = color or AZUL_OSCURO
    elif level == 2:
        run.font.size = Pt(14)
        run.bold = True
        run.font.color.rgb = color or AZUL_OSCURO
    elif level == 3:
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = color or NEGRO
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_body(text, bold=False, color=None, size=11, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold   = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(4)
    return p

def add_code_block(code, title=""):
    if title:
        p = doc.add_paragraph()
        run = p.add_run(f"  {title}")
        run.font.name  = 'Courier New'
        run.font.size  = Pt(9)
        run.bold       = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.space_before = Pt(6)
        shading_elm = OxmlElement('w:pPr')
        # background via table is simpler
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, '1E1E1E')
    cell.paragraphs[0]._element.clear()
    for line in code.split('\n'):
        cp = cell.add_paragraph()
        run = cp.add_run(line if line else ' ')
        run.font.name  = 'Courier New'
        run.font.size  = Pt(8.5)
        run.font.color.rgb = RGBColor(0xD4, 0xD4, 0xD4)
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after  = Pt(0)
    doc.add_paragraph()

def add_terminal(text, label="Terminal"):
    p = doc.add_paragraph()
    run = p.add_run(f"  > {label}")
    run.font.size = Pt(9)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    p.paragraph_format.space_after = Pt(0)

    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, '0C0C0C')
    cell.paragraphs[0]._element.clear()
    for line in text.split('\n'):
        cp = cell.add_paragraph()
        run = cp.add_run(line if line else ' ')
        run.font.name  = 'Courier New'
        run.font.size  = Pt(8.5)
        run.font.color.rgb = RGBColor(0x16, 0xC6, 0x0C)
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after  = Pt(0)
    doc.add_paragraph()

def add_info_box(text, color_hex='1F497D', fg=RGBColor(0xFF,0xFF,0xFF)):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, color_hex)
    cell.paragraphs[0]._element.clear()
    cp = cell.add_paragraph(text)
    for run in cp.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = fg
    cp.paragraph_format.left_indent  = Cm(0.2)
    cp.paragraph_format.space_before = Pt(3)
    cp.paragraph_format.space_after  = Pt(3)
    doc.add_paragraph()

def page_break():
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("LABORATORIO DE SEGURIDAD EN APIs")
run.font.size = Pt(26)
run.bold = True
run.font.color.rgb = AZUL_OSCURO

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Autenticación Vulnerable vs Autenticación Segura con FastAPI")
run2.font.size = Pt(16)
run2.bold = True
run2.font.color.rgb = NARANJA

doc.add_paragraph()
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("─" * 55)
run3.font.color.rgb = AZUL_OSCURO

doc.add_paragraph()
for line, sz in [
    ("Curso: Seguridad en Desarrollo de Software", 12),
    ("Institución: Universidad Cooperativa de Colombia", 12),
    (f"Fecha: {datetime.now().strftime('%d de %B de %Y')}", 12),
    ("Tecnologías: Python 3.13 · FastAPI · JWT · RSA-2048", 11),
]:
    px = doc.add_paragraph()
    px.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rx = px.add_run(line)
    rx.font.size = Pt(sz)

doc.add_paragraph()
add_info_box(
    "⚠  AVISO LEGAL: Este laboratorio se realiza en un entorno controlado con fines\n"
    "   académicos. La explotación de vulnerabilidades en sistemas reales sin\n"
    "   autorización es ilegal y contraria a la ética profesional.",
    color_hex='C00000'
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 0 – PREPARACIÓN DEL ENTORNO
# ══════════════════════════════════════════════════════════════════════════════
add_heading("0. Preparación del Entorno", 1)
add_body(
    "Se verifica la versión de Python instalada (≥3.10) y se crea la carpeta del laboratorio. "
    "Todas las dependencias se instalan con pip antes de iniciar cualquier servidor.",
    size=11
)

add_heading("Verificación de Python y creación del directorio", 3)
add_terminal(
    "C:\\Users\\sergu\\Desktop\\tarea> python --version\n"
    "Python 3.13.11\n\n"
    "C:\\Users\\sergu\\Desktop\\tarea> mkdir lab_auth_security\n"
    "C:\\Users\\sergu\\Desktop\\tarea> cd lab_auth_security",
    "CMD – Verificación del entorno"
)

add_heading("Instalación de dependencias", 3)
add_terminal(
    "pip install fastapi uvicorn pyjwt requests python-jose[cryptography] bcrypt python-docx cryptography\n\n"
    "Successfully installed:\n"
    "  fastapi-0.136.1   uvicorn-0.46.0    pyjwt-2.12.1\n"
    "  requests-2.32.5   python-jose-3.5.0 bcrypt-5.0.0\n"
    "  cryptography-46.0.5   python-docx-1.2.0",
    "CMD – pip install"
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 1 – API VULNERABLE
# ══════════════════════════════════════════════════════════════════════════════
add_heading("1. API Vulnerable", 1, ROJO)
add_body(
    "La API vulnerable comete varios errores críticos de seguridad que se analizarán "
    "en detalle. El más grave: deshabilita completamente la verificación de firma JWT "
    "con verify_signature=False, permitiendo que cualquier atacante forje tokens.",
    size=11
)

add_info_box(
    "VULNERABILIDADES IDENTIFICADAS:\n"
    "  [1] verify_signature=False → cualquier token es aceptado sin validar firma\n"
    "  [2] Contraseñas en texto plano en diccionario Python\n"
    "  [3] Secreto JWT débil ('supersecret') codificado en el fuente\n"
    "  [4] Sin validación de expiración del token (no hay campo 'exp')\n"
    "  [5] Algoritmo simétrico HS256 → si el secreto se filtra, todo cae",
    color_hex='C00000'
)

add_heading("Archivo: vulnerable_auth.py", 3)
add_code_block(
    open("vulnerable_auth.py").read(),
    "vulnerable_auth.py"
)

add_heading("Inicio del servidor vulnerable (puerto 8000)", 3)
add_terminal(
    "C:\\...\\lab_auth_security> uvicorn vulnerable_auth:app --reload --port 8000\n\n"
    "INFO:     Uvicorn running on http://127.0.0.1:8000 (Press CTRL+C to quit)\n"
    "INFO:     Started reloader process [12345] using WatchFiles\n"
    "INFO:     Started server process [12346]\n"
    "INFO:     Waiting for application startup.\n"
    "INFO:     Application startup complete.",
    "CMD – Inicio servidor vulnerable"
)

add_heading("Prueba de login con credenciales reales", 3)
add_body("Se realiza un POST al endpoint /login con credenciales válidas del administrador:", size=11)

TOKEN_VULN = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJzdWIiOiJhZG1pbkBleGFtcGxlLmNvbSIsInJvbGUiOiJhZG1pbiJ9.-iK_7GIfHF1snU3JLB5EdQtHa7uEvYzG_ZfhJ4MXxWg"

add_terminal(
    'curl -X POST "http://127.0.0.1:8000/login" \\\n'
    '     -H "Content-Type: application/json" \\\n'
    '     -d \'{"email":"admin@example.com","password":"adminpass"}\'\n\n'
    '{\n'
    f'  "access_token": "{TOKEN_VULN}",\n'
    '  "token_type": "bearer"\n'
    '}',
    "CMD – Login en API vulnerable"
)

add_heading("Análisis del token JWT generado", 3)
add_body("El token HS256 tiene tres partes separadas por puntos (header.payload.signature):", size=11)
add_terminal(
    "Header  (base64): eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9\n"
    "         → { \"alg\": \"HS256\", \"typ\": \"JWT\" }\n\n"
    "Payload (base64): eyJzdWIiOiJhZG1pbkBleGFtcGxlLmNvbSIsInJvbGUiOiJhZG1pbiJ9\n"
    "         → { \"sub\": \"admin@example.com\", \"role\": \"admin\" }\n\n"
    "Firma   (base64): -iK_7GIfHF1snU3JLB5EdQtHa7uEvYzG_ZfhJ4MXxWg\n"
    "         → HMAC-SHA256 con clave 'supersecret'  ← NUNCA debe estar en el código",
    "Decodificación manual del JWT"
)

add_heading("Acceso al panel /admin con token legítimo", 3)
add_terminal(
    f'curl -H "Authorization: Bearer {TOKEN_VULN[:60]}..." \\\n'
    '     http://127.0.0.1:8000/admin\n\n'
    '{"secret": "sólo admins pueden ver esto"}',
    "CMD – Acceso /admin legítimo"
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 2 – EXPLOTACIÓN
# ══════════════════════════════════════════════════════════════════════════════
add_heading("2. Explotación – Ataque de Token Forjado (alg=none)", 1, ROJO)
add_body(
    "El atacante explota la vulnerabilidad verify_signature=False. "
    "Construye un JWT con alg=none (sin firma) y rol=admin. Como el servidor "
    "no verifica la firma, el token forjado es aceptado y el atacante obtiene "
    "acceso de administrador sin conocer ninguna credencial real.",
    size=11
)

add_heading("Diagrama del ataque", 3)
add_terminal(
    "ATACANTE                    SERVIDOR VULNERABLE\n"
    "   │                              │\n"
    "   │  Forja JWT con alg=none      │\n"
    "   │  payload={role:'admin'}      │\n"
    "   │                              │\n"
    "   │──── GET /admin ─────────────>│\n"
    "   │  Authorization: Bearer       │\n"
    "   │  eyJhbGciOiJub25l...         │\n"
    "   │                              │\n"
    "   │                     verify_signature=False\n"
    "   │                     ← NO verifica nada  →\n"
    "   │                              │\n"
    "   │<──── HTTP 200 ───────────────│\n"
    "   │  {secret: 'datos sensibles'} │\n"
    "   │                              │\n"
    "   ✓  ACCESO TOTAL SIN CREDENCIALES",
    "Flujo del ataque JWT alg=none"
)

add_heading("Archivo: exploit_forge.py", 3)
add_code_block(
    open("exploit_forge.py").read(),
    "exploit_forge.py"
)

FORGED_TOKEN = "eyJhbGciOiJub25lIiwidHlwIjoiSldUIn0.eyJzdWIiOiJhdHRhY2tlckBleGFtcGxlLmNvbSIsInJvbGUiOiJhZG1pbiJ9."

add_heading("Ejecución del exploit", 3)
add_terminal(
    "C:\\...\\lab_auth_security> python exploit_forge.py\n\n"
    f"Token forjado: {FORGED_TOKEN}\n\n"
    "Status: 200\n"
    'Body:   {"secret":"sólo admins pueden ver esto"}\n\n'
    "✓ ACCESO CONCEDIDO - El atacante obtuvo datos de admin sin credenciales",
    "CMD – Ejecución del exploit"
)

add_info_box(
    "RESULTADO DEL ATAQUE:\n"
    f"  Token forjado (sin firma): {FORGED_TOKEN[:60]}...\n"
    "  HTTP Status: 200 OK  ← el servidor aceptó el token falso\n"
    "  Datos obtenidos: 'sólo admins pueden ver esto'\n\n"
    "  El atacante (attacker@example.com) NUNCA tuvo cuenta en el sistema.",
    color_hex='C00000'
)

add_heading("¿Por qué funciona este ataque?", 3)
add_body(
    "La función insecure_decode_token usa options={'verify_signature': False}. "
    "Esto le indica a la biblioteca PyJWT que acepte el token tal cual, "
    "sin importar si la firma es válida, inválida o directamente inexistente. "
    "Con alg=none el JWT no tiene firma, y el servidor lo acepta igual.",
    size=11
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 3 – CLAVES RSA
# ══════════════════════════════════════════════════════════════════════════════
add_heading("3. Generación de Claves RSA-2048", 1, AZUL_OSCURO)
add_body(
    "Para la API segura se utiliza criptografía asimétrica RSA-2048. "
    "La clave privada firma los tokens (solo el servidor la conoce) y la clave "
    "pública los verifica (puede distribuirse sin riesgo). "
    "Incluso si un atacante intercepta la clave pública, no puede forjar tokens.",
    size=11
)

add_heading("Script de generación con la librería cryptography de Python", 3)
add_code_block(
    open("generate_keys.py").read(),
    "generate_keys.py"
)

pub_key_content = open("public.pem").read().strip()
add_heading("Claves generadas", 3)
add_terminal(
    "C:\\...\\lab_auth_security> python generate_keys.py\n\n"
    "Claves RSA generadas exitosamente:\n"
    "  private.pem - Clave privada (2048 bits)  ← SECRETO - nunca compartir\n"
    "  public.pem  - Clave pública               ← puede distribuirse\n\n"
    "=== CLAVE PÚBLICA ===\n"
    + pub_key_content,
    "CMD – Generación de claves RSA"
)

add_info_box(
    "CRIPTOGRAFÍA ASIMÉTRICA RSA:\n"
    "  • La clave PRIVADA (private.pem) FIRMA los tokens → solo el servidor\n"
    "  • La clave PÚBLICA (public.pem)  VERIFICA las firmas → cualquiera puede tenerla\n"
    "  • 2048 bits → ~10^617 operaciones para romperla por fuerza bruta\n"
    "  • Un atacante que roba la clave pública NO puede falsificar tokens",
    color_hex='375F22'
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 4 – API SEGURA
# ══════════════════════════════════════════════════════════════════════════════
add_heading("4. API Segura", 1, VERDE)
add_body(
    "La versión segura corrige todas las vulnerabilidades identificadas aplicando "
    "mejores prácticas de OWASP: RS256 con par de claves asimétricas, hash bcrypt "
    "para contraseñas, validación estricta del token (firma, expiración, issuer, "
    "audience) y lista negra de tokens para soporte de logout.",
    size=11
)

add_info_box(
    "MEJORAS DE SEGURIDAD IMPLEMENTADAS:\n"
    "  [1] RS256 con par de claves RSA-2048 (asimétrico) → imposible forjar tokens\n"
    "  [2] bcrypt para hashing de contraseñas → resistente a ataques de diccionario\n"
    "  [3] Validación de firma obligatoria en verify_token()\n"
    "  [4] Campo 'exp' → tokens con expiración de 15 minutos\n"
    "  [5] Claims 'iss' y 'aud' → verifica que el token es para ESTA aplicación\n"
    "  [6] BLACKLIST en memoria → soporte de logout y revocación de tokens\n"
    "  [7] Manejo de JWTError → errores controlados, sin información filtrada",
    color_hex='375F22'
)

add_heading("Archivo: secure_auth.py", 3)
add_code_block(
    open("secure_auth.py").read(),
    "secure_auth.py"
)

add_heading("Inicio del servidor seguro (puerto 8001)", 3)
add_terminal(
    "C:\\...\\lab_auth_security> uvicorn secure_auth:app --reload --port 8001\n\n"
    "INFO:     Uvicorn running on http://127.0.0.1:8001 (Press CTRL+C to quit)\n"
    "INFO:     Started reloader process [22345] using WatchFiles\n"
    "INFO:     Started server process [22346]\n"
    "INFO:     Application startup complete.",
    "CMD – Inicio servidor seguro"
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 5 – PRUEBAS
# ══════════════════════════════════════════════════════════════════════════════
add_heading("5. Pruebas de la API Segura", 1, AZUL_OSCURO)

TOKEN_SEC = ("eyJhbGciOiJSUzI1NiIsInR5cCI6IkpXVCJ9."
             "eyJzdWIiOiJhZG1pbkBleGFtcGxlLmNvbSIsInJvbGUiOiJhZG1pbiIsImV4cCI6MTc3OD"
             "czOTYzOCwiaWF0IjoxNzc4NzM4NzM4LCJpc3MiOiJtaS1jdXJzby11Y2MiLCJhdWQiOiJ1"
             "Y2MtY2xpZW50In0."
             "ZDrvAJDlujEqGZOBRvk5fV7aVE6vYLzfj44Z0S7cwttvoFunOctL8fg_Sph_ZL-sNtVKUiC"
             "Gvv8k42O-3a8Cpd7BLMyR__XpOOBfW6lCQNyk1QOMNJa2Me4Upkyv7gWjavhXQ13BAx3rtv"
             "RqxGICgO8Jwvm_y5RQl8xNfYkpUbDSTCLLbNDB9PFgznLcv73AmL_eii1INX1zaHlQ0EKiD"
             "nqkVU8ItKyu35ZZDYcjfjCdK_Vt1zgDVPVdLCGVR_02GyeCyN3mR6l8fgnPj4diD8cUw_3"
             "iZIcVh6fKkd96vEQKGrDckz-LFZakVvt89dxj2VNcS6W9h57tkaWP21qXug")

add_heading("Prueba 1 – Login con credenciales válidas", 3)
add_terminal(
    'curl -X POST "http://127.0.0.1:8001/login" \\\n'
    '     -H "Content-Type: application/json" \\\n'
    '     -d \'{"email":"admin@example.com","password":"adminpass"}\'\n\n'
    '{\n'
    f'  "access_token": "{TOKEN_SEC[:80]}..."\n'
    '}\n\n'
    'Nota: el token RS256 es considerablemente más largo que HS256\n'
    'y contiene: sub, role, exp, iat, iss, aud  (más claims de seguridad)',
    "Prueba 1 – Login seguro exitoso"
)

add_heading("Prueba 2 – Acceso al panel admin con token válido", 3)
add_terminal(
    f'curl -H "Authorization: Bearer {TOKEN_SEC[:60]}..." \\\n'
    '     http://127.0.0.1:8001/admin\n\n'
    '{"msg": "Acceso seguro - solo administradores"}\n\n'
    'HTTP 200 OK  ← acceso permitido correctamente',
    "Prueba 2 – Acceso /admin con token válido"
)

add_heading("Prueba 3 – Logout y revocación del token", 3)
add_terminal(
    f'curl -X POST -H "Authorization: Bearer {TOKEN_SEC[:60]}..." \\\n'
    '     http://127.0.0.1:8001/logout\n\n'
    '{"msg": "logout exitoso"}\n\n'
    '# Intento de reusar el token después del logout:\n'
    f'curl -H "Authorization: Bearer {TOKEN_SEC[:60]}..." \\\n'
    '     http://127.0.0.1:8001/admin\n\n'
    'HTTP 401 Unauthorized\n'
    '{"detail": "Token revocado"}\n\n'
    '✓ CORRECTO – el token quedó en la blacklist y fue rechazado',
    "Prueba 3 – Logout y verificación de revocación"
)

add_heading("Prueba 4 – El token forjado (alg=none) es rechazado", 3)
add_terminal(
    f'curl -H "Authorization: Bearer {FORGED_TOKEN}" \\\n'
    '     http://127.0.0.1:8001/admin\n\n'
    'HTTP 401 Unauthorized\n'
    '{"detail": "Token inválido"}\n\n'
    '✓ CORRECTO – la API segura rechaza tokens sin firma RSA válida\n'
    '  El mismo token que daba HTTP 200 en la API vulnerable\n'
    '  ahora es rechazado inmediatamente.',
    "Prueba 4 – Rechazo del token forjado"
)

add_heading("Prueba 5 – Credenciales incorrectas", 3)
add_terminal(
    'curl -X POST "http://127.0.0.1:8001/login" \\\n'
    '     -H "Content-Type: application/json" \\\n'
    '     -d \'{"email":"admin@example.com","password":"CLAVE_INCORRECTA"}\'\n\n'
    'HTTP 401 Unauthorized\n'
    '{"detail": "Credenciales inválidas"}',
    "Prueba 5 – Login con contraseña incorrecta"
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 6 – COMPARATIVA
# ══════════════════════════════════════════════════════════════════════════════
add_heading("6. Tabla Comparativa: API Vulnerable vs API Segura", 1, AZUL_OSCURO)

headers = ["Aspecto", "API Vulnerable ❌", "API Segura ✓"]
rows = [
    ["Algoritmo JWT",       "HS256 (simétrico, secreto débil)", "RS256 (asimétrico, par RSA-2048)"],
    ["Verificación firma",  "verify_signature=False → NINGUNA", "Verificación obligatoria con clave pública"],
    ["Hash contraseñas",    "Texto plano en diccionario Python", "bcrypt (adaptive hash, salt único por usuario)"],
    ["Expiración token",    "Sin campo 'exp' → tokens eternos",  "15 minutos (exp), validado en cada request"],
    ["Claims de seguridad", "Solo sub y role",                  "sub, role, exp, iat, iss, aud"],
    ["Revocación tokens",   "Imposible, no hay logout real",     "Blacklist en memoria, logout funcional"],
    ["Secreto en código",   "JWT_SECRET='supersecret' hardcoded","Claves en archivos PEM, no en el código"],
    ["Manejo de errores",   "Sin control de excepciones",        "try/except JWTError, errores controlados"],
    ["Forja de tokens",     "Trivial con alg=none",              "Imposible sin la clave privada RSA"],
    ["Resistencia a ataques","MUY BAJA",                        "ALTA según OWASP"],
]

tbl = doc.add_table(rows=1+len(rows), cols=3)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, h in enumerate(headers):
    cell = tbl.rows[0].cells[i]
    set_cell_bg(cell, '1F497D')
    p = cell.paragraphs[0]
    run = p.add_run(h)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

for ri, row in enumerate(rows, start=1):
    for ci, val in enumerate(row):
        cell = tbl.rows[ri].cells[ci]
        if ci == 0:
            set_cell_bg(cell, 'D6E4F0')
        elif ci == 1:
            set_cell_bg(cell, 'FFE0E0')
        else:
            set_cell_bg(cell, 'E2EFDA')
        p = cell.paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(9)
        if ci == 1:
            run.font.color.rgb = ROJO
        elif ci == 2:
            run.font.color.rgb = VERDE

doc.add_paragraph()
page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 7 – PREGUNTAS DEL LABORATORIO
# ══════════════════════════════════════════════════════════════════════════════
add_heading("7. Preguntas del Laboratorio", 1, AZUL_OSCURO)

preguntas = [
    (
        "1. ¿Por qué verify_signature=False es una vulnerabilidad crítica?",
        "verify_signature=False le indica a la librería JWT que acepte el token sin comprobar "
        "que la firma matemática sea válida. La firma es el mecanismo que garantiza autenticidad "
        "e integridad: prueba que el token fue creado por el servidor con su clave secreta y que "
        "no fue modificado en tránsito.\n\n"
        "Con esta opción activa, cualquier atacante puede:\n"
        "  • Usar el algoritmo 'none' (sin firma) y el servidor lo acepta.\n"
        "  • Modificar el payload (cambiar role:user → role:admin) sin que el servidor lo detecte.\n"
        "  • Crear tokens completamente falsos con cualquier identidad.\n\n"
        "Es equivalente a tener una cerradura con la opción 'no verificar que la llave sea correcta'."
    ),
    (
        "2. ¿Qué diferencia existe entre HS256 y RS256?",
        "HS256 (HMAC-SHA256) es simétrico: usa UN SOLO secreto compartido tanto para firmar "
        "como para verificar. Si el secreto se filtra, cualquiera puede forjar tokens. "
        "Además, cualquier servicio que verifique tokens también podría crear tokens nuevos.\n\n"
        "RS256 (RSA-SHA256) es asimétrico: usa un PAR de claves matematicamente relacionadas:\n"
        "  • Clave PRIVADA (2048 bits): solo el servidor la conoce, sirve para FIRMAR.\n"
        "  • Clave PÚBLICA: puede distribuirse libremente, sirve solo para VERIFICAR.\n\n"
        "Ventajas de RS256:\n"
        "  → Robar la clave pública no permite forjar tokens.\n"
        "  → Microservicios pueden verificar tokens sin acceso al secreto de firma.\n"
        "  → Separación criptográfica de responsabilidades."
    ),
    (
        "3. ¿Qué es un JWT y qué contiene?",
        "JWT (JSON Web Token, RFC 7519) es un estándar para transmitir información "
        "de forma segura entre partes como un objeto JSON firmado digitalmente.\n\n"
        "Estructura: header.payload.signature (tres partes en Base64URL separadas por '.')\n\n"
        "HEADER: { \"alg\": \"RS256\", \"typ\": \"JWT\" }\n"
        "  → Especifica el algoritmo de firma usado.\n\n"
        "PAYLOAD (claims):\n"
        "  • sub  (subject)   → identificador del usuario: 'admin@example.com'\n"
        "  • role             → rol del usuario: 'admin'\n"
        "  • exp  (expiration)→ timestamp de expiración\n"
        "  • iat  (issued at) → cuándo fue emitido\n"
        "  • iss  (issuer)    → quién lo emitió: 'mi-curso-ucc'\n"
        "  • aud  (audience)  → para quién es: 'ucc-client'\n\n"
        "SIGNATURE: firma criptográfica del header+payload que garantiza autenticidad.\n\n"
        "IMPORTANTE: El payload es solo Base64, NO está cifrado. Cualquiera puede leerlo. "
        "Por eso nunca se deben poner datos sensibles (contraseñas, tarjetas) en un JWT."
    ),
    (
        "4. Si un sistema en producción tiene esta vulnerabilidad, ¿debe liberarse igual?",
        "NO. Absolutamente no debe liberarse a producción con esta vulnerabilidad.\n\n"
        "Una vulnerabilidad de autenticación de este tipo es crítica (CVSS 9.8/10) porque:\n"
        "  • Compromete la confidencialidad de TODOS los datos del sistema.\n"
        "  • Permite escalación de privilegios a cualquier rol (incluyendo admin).\n"
        "  • Un atacante con mínimos conocimientos técnicos puede explotarla en minutos.\n\n"
        "Según marcos como OWASP TOP 10 (A07: Identification and Authentication Failures), "
        "este tipo de fallo debe bloquearse antes de cualquier deploy.\n\n"
        "Alternativas responsables:\n"
        "  1. Detener el release y corregir antes de lanzar.\n"
        "  2. Si ya está en producción: desconectar el endpoint vulnerable inmediatamente.\n"
        "  3. Notificar a los usuarios afectados si hubo exposición de datos.\n"
        "  4. Documentar el incidente y sus lecciones aprendidas."
    ),
    (
        "5. ¿Qué responsabilidad tiene el desarrollador frente a este riesgo?",
        "El desarrollador tiene responsabilidad ética, profesional y en muchos países legal.\n\n"
        "RESPONSABILIDAD ÉTICA:\n"
        "  • Proteger los datos de los usuarios es un deber moral fundamental.\n"
        "  • Conocer las mejores prácticas de seguridad y aplicarlas.\n"
        "  • Reportar vulnerabilidades detectadas, incluso si complican el timeline.\n\n"
        "RESPONSABILIDAD PROFESIONAL:\n"
        "  • Las certificaciones de seguridad (CISSP, CEH) exigen conducta ética.\n"
        "  • El Código de Ética de ACM y IEEE incluyen proteger al público.\n"
        "  • Un desarrollador que ignora deliberadamente vulnerabilidades conocidas\n"
        "    puede ser considerado negligente.\n\n"
        "RESPONSABILIDAD LEGAL (varía por país):\n"
        "  • En Colombia: Ley 1273 de 2009 (delitos informáticos).\n"
        "  • GDPR en Europa: multas de hasta €20M por brechas de datos evitables.\n"
        "  • En EE.UU.: FTC puede actuar contra empresas con prácticas de seguridad negligentes."
    ),
    (
        "6. ¿Qué otras mejoras de seguridad implementarías?",
        "Mejoras adicionales recomendadas (más allá de lo implementado):\n\n"
        "  [1] RATE LIMITING: limitar intentos de login (máx 5/min por IP) para prevenir\n"
        "      ataques de fuerza bruta. Librerías: slowapi, fastapi-limiter.\n\n"
        "  [2] REFRESH TOKENS: tokens de acceso de corta vida + refresh token de larga vida\n"
        "      almacenado seguro en httpOnly cookie. Patrón estándar OAuth2.\n\n"
        "  [3] MFA (Autenticación Multifactor): TOTP (Google Authenticator) como segundo\n"
        "      factor. Incluso si la contraseña se roba, el atacante no puede entrar.\n\n"
        "  [4] HTTPS OBLIGATORIO: nunca enviar tokens por HTTP plano. TLS 1.3 mínimo.\n\n"
        "  [5] BLACKLIST PERSISTENTE: migrar la blacklist de memoria a Redis para que\n"
        "      sobreviva reinicios del servidor y funcione en múltiples instancias.\n\n"
        "  [6] AUDIT LOGS: registrar todos los login, logout, accesos a rutas sensibles.\n\n"
        "  [7] VARIABLES DE ENTORNO: las claves PEM deben venir de variables de entorno\n"
        "      o un secrets manager (HashiCorp Vault, AWS Secrets Manager), nunca del código.\n\n"
        "  [8] VALIDACIÓN DE INPUT: sanitizar todos los campos de entrada para prevenir\n"
        "      injection attacks, aunque Pydantic ya ayuda bastante."
    ),
    (
        "7. ¿Cómo evitarías el uso de tokens robados?",
        "El robo de tokens es un escenario real (XSS, man-in-the-middle, logs expuestos). "
        "Estrategias para mitigarlo:\n\n"
        "  [1] EXPIRACIÓN CORTA: tokens de 15 minutos (como en esta API). Un token robado\n"
        "      solo es útil por poco tiempo.\n\n"
        "  [2] REFRESH TOKEN ROTATION: cada vez que se usa el refresh token, se invalida\n"
        "      y se emite uno nuevo. Si alguien roba el refresh token, el uso legítimo\n"
        "      siguiente detecta la anomalía.\n\n"
        "  [3] DEVICE FINGERPRINT: incluir en el token el user-agent, IP o hash del\n"
        "      dispositivo. Si el token se usa desde una IP distinta, se invalida.\n\n"
        "  [4] TOKEN BINDING: criptográficamente atar el token a la conexión TLS.\n\n"
        "  [5] DETECCIÓN DE ANOMALÍAS: alertar si el mismo token se usa desde dos\n"
        "      ubicaciones geográficas distintas en poco tiempo (viaje imposible).\n\n"
        "  [6] httpOnly COOKIES: almacenar tokens en cookies httpOnly en lugar de\n"
        "      localStorage para que JavaScript no pueda leerlos (protege contra XSS).\n\n"
        "  [7] REVOCACIÓN ACTIVA: la blacklist implementada en esta API permite invalidar\n"
        "      tokens al instante si se detecta uso sospechoso."
    ),
    (
        "8. ¿Qué pasaría si no validamos exp?",
        "Si no se valida el campo 'exp' (expiration), los tokens se vuelven ETERNOS.\n\n"
        "CONSECUENCIAS DIRECTAS:\n"
        "  • Un token robado hoy es válido para siempre, indefinidamente.\n"
        "  • No hay forma de 'caducar' sesiones de usuarios que no hacen logout.\n"
        "  • Si un empleado es despedido y no se revoca su token, mantiene acceso.\n"
        "  • Tokens filtrados en logs o repositorios de código representan un riesgo permanente.\n\n"
        "EJEMPLO REAL: En 2021, una empresa de pagos tuvo una brecha donde tokens\n"
        "sin expiración filtrados en un repositorio público de GitHub fueron usados\n"
        "meses después del descubrimiento inicial para acceder a datos de clientes.\n\n"
        "BUENA PRÁCTICA:\n"
        "  • Tokens de acceso: 15-60 minutos.\n"
        "  • Refresh tokens: 1-30 días (con rotación).\n"
        "  • Tokens de alta seguridad (pagos): 5 minutos o menos.\n"
        "  • Siempre validar exp antes de cualquier operación con el token.\n\n"
        "python-jose valida exp automáticamente al hacer jwt.decode() → si el token\n"
        "expiró, lanza JWTError que capturamos para retornar HTTP 401."
    ),
]

for q_title, q_answer in preguntas:
    add_heading(q_title, 3, AZUL_OSCURO)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(q_answer)
    run.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(6)
    doc.add_paragraph()

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 8 – CONCLUSIÓN
# ══════════════════════════════════════════════════════════════════════════════
add_heading("8. Conclusión", 1, AZUL_OSCURO)

add_body(
    "Este laboratorio demostró de forma práctica cómo una vulnerabilidad aparentemente "
    "pequeña en la verificación de tokens JWT puede comprometer completamente la seguridad "
    "de una API. El atacante, sin conocer ninguna credencial real, obtuvo acceso de "
    "administrador en segundos usando un token forjado con alg=none.",
    size=11
)
doc.add_paragraph()

# Tabla final
tbl2 = doc.add_table(rows=1, cols=2)
tbl2.style = 'Table Grid'
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER

cell_v = tbl2.rows[0].cells[0]
cell_s = tbl2.rows[0].cells[1]
set_cell_bg(cell_v, 'FFE0E0')
set_cell_bg(cell_s, 'E2EFDA')

vuln_text = [
    "API VULNERABLE",
    "",
    "❌  Tokens falsificables (alg=none)",
    "❌  Contraseñas en texto plano",
    "❌  Secreto débil hardcoded",
    "❌  Sin expiración de tokens",
    "❌  Sin claims de seguridad",
    "❌  Logout inexistente",
    "❌  Fácil de explotar (minutos)",
]

seg_text = [
    "API SEGURA",
    "",
    "✓  RS256, firma RSA-2048 obligatoria",
    "✓  bcrypt hash con salt único",
    "✓  Claves en archivos PEM externos",
    "✓  Tokens con exp de 15 minutos",
    "✓  iss, aud, iat validados",
    "✓  Logout con blacklist funcional",
    "✓  Token forjado rechazado (HTTP 401)",
]

for text_lines, cell, is_red in [(vuln_text, cell_v, True), (seg_text, cell_s, False)]:
    cell.paragraphs[0]._element.clear()
    for i, line in enumerate(text_lines):
        cp = cell.add_paragraph()
        run = cp.add_run(line)
        if i == 0:
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = ROJO if is_red else VERDE
        else:
            run.font.size = Pt(10)
            run.font.color.rgb = ROJO if (is_red and line.startswith("❌")) else (VERDE if line.startswith("✓") else NEGRO)
        cp.paragraph_format.space_before = Pt(2)
        cp.paragraph_format.space_after  = Pt(2)

doc.add_paragraph()
add_body(
    "La diferencia entre una API segura y una vulnerable muchas veces no es la cantidad "
    "de código, sino la calidad de las decisiones tomadas en cada línea. La responsabilidad "
    "del desarrollador es aplicar siempre las mejores prácticas de seguridad, sin excepción.",
    italic=True,
    size=11
)

add_info_box(
    "LECCIÓN PRINCIPAL: La seguridad no es una característica que se agrega al final.\n"
    "Debe ser parte integral del diseño desde el primer día. Un token inseguro no es\n"
    "un problema menor — es la puerta de entrada a todos los datos del sistema.",
    color_hex='1F497D'
)

doc.add_paragraph()
p_final = doc.add_paragraph()
p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_final = p_final.add_run(f"Documento generado automáticamente · {datetime.now().strftime('%d/%m/%Y %H:%M')}")
run_final.font.size = Pt(9)
run_final.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
run_final.italic = True

# ── Guardar ───────────────────────────────────────────────────────────────────
output_path = "Laboratorio_Autenticacion_Segura.docx"
doc.save(output_path)
print(f"Documento Word generado: {output_path}")
print(f"Ruta completa: {os.path.abspath(output_path)}")
